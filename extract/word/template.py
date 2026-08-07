"""Génère extract/templates/Editions_Particulieres.dotx — styles Manuel V1."""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, Twips

from . import styles as S
from .theme import (
    ACCENT,
    BORDER,
    CALLOUT_BORDER_SHADE,
    CALLOUT_BORDER_SPACE,
    CALLOUT_BORDER_SZ,
    CALLOUT_BORDER_THEME,
    FONT_BODY,
    FONT_BODY_FALLBACK,
    FONT_DISPLAY,
    FONT_DISPLAY_FALLBACK,
    INK,
    LINE_BODY,
    LINE_TIGHT,
    MUTED,
    QUOTE_BG,
    SIZE_BODY,
    SIZE_H1,
    SIZE_H2,
    SIZE_H3,
    SIZE_H4,
    SIZE_H5,
    SIZE_H6,
    SIZE_SUBTITLE,
)

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "templates" / "Editions_Particulieres.dotx"


def _set_font(style, name: str, fallback: str, size_pt: float, *, bold=None, italic=None, color=None):
    font = style.font
    font.name = name
    font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if color is not None:
        font.color.rgb = color
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" '
            f'w:cs="{fallback}" w:eastAsia="{fallback}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:cs"), fallback)


def _spacing(style, before=0, after=6, line=LINE_BODY):
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def _align(style, alignment):
    style.paragraph_format.alignment = alignment


def _outline(style, level: int | None):
    pPr = style.element.get_or_add_pPr()
    for old in pPr.findall(qn("w:outlineLvl")):
        pPr.remove(old)
    if level is not None:
        pPr.append(parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="{int(level)}"/>'))


def _shading(style, fill_hex: str):
    pPr = style.element.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    pPr.append(
        parse_xml(
            f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>'
        )
    )


def _left_border(style, color_hex: str = "C4A35A", sz: int = 16, space: int = 8):
    pPr = style.element.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pPr.append(
        parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="{space}" w:color="{color_hex}"/>'
            f"</w:pBdr>"
        )
    )


def _theme_left_border(
    style,
    *,
    sz: int = CALLOUT_BORDER_SZ,
    space: int = CALLOUT_BORDER_SPACE,
) -> None:
    """Barre gauche thème Word (Accent 1 Darker 25 %)."""
    pPr = style.element.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pPr.append(
        parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:left w:val="single" w:sz="{sz}" w:space="{space}" w:color="auto" '
            f'w:themeColor="{CALLOUT_BORDER_THEME}" w:themeShade="{CALLOUT_BORDER_SHADE}"/>'
            f"</w:pBdr>"
        )
    )


def _bottom_rule(style, color_hex: str = "C4A35A", sz: int = 8):
    pPr = style.element.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pPr.append(
        parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="4" w:color="{color_hex}"/>'
            f"</w:pBdr>"
        )
    )


def _get_or_add(doc: Document, name: str, style_type):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_font(normal, FONT_BODY, FONT_BODY_FALLBACK, SIZE_BODY, color=INK)
    _spacing(normal, after=8, line=LINE_BODY)
    _align(normal, WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Title : non utilisé — laisser le défaut Word
    title = doc.styles["Title"]
    _set_font(title, FONT_DISPLAY, FONT_DISPLAY_FALLBACK, SIZE_H1, color=INK)
    _outline(title, None)

    # Subtitle — relations sous le titre (.site-crumb / kicker)
    sub = _get_or_add(doc, S.STYLE_SUBTITLE, WD_STYLE_TYPE.PARAGRAPH)
    sub.base_style = normal
    _set_font(sub, FONT_BODY, FONT_BODY_FALLBACK, SIZE_SUBTITLE, color=MUTED)
    _spacing(sub, before=0, after=10, line=LINE_TIGHT)
    sub.font.all_caps = False
    sub.font.small_caps = True
    _outline(sub, None)

    # Heading 1 — titre de page Notion (.site-title)
    h1 = doc.styles[S.STYLE_H1]
    h1.base_style = normal
    _set_font(h1, FONT_DISPLAY, FONT_DISPLAY_FALLBACK, SIZE_H1, bold=True, color=INK)
    _spacing(h1, before=0, after=14, line=LINE_TIGHT)
    _bottom_rule(h1)
    _outline(h1, 0)

    # Heading 2 — Notion heading_1 (.legal-prose h2)
    h2 = doc.styles[S.STYLE_H2]
    h2.base_style = normal
    _set_font(h2, FONT_DISPLAY, FONT_DISPLAY_FALLBACK, SIZE_H2, bold=True, color=INK)
    _spacing(h2, before=18, after=8, line=LINE_TIGHT)
    h2.paragraph_format.keep_with_next = True
    _outline(h2, 1)

    # Heading 3 — Notion heading_2 (.legal-prose h3)
    h3 = doc.styles[S.STYLE_H3]
    h3.base_style = normal
    _set_font(h3, FONT_BODY, FONT_BODY_FALLBACK, SIZE_H3, bold=True, color=INK)
    _spacing(h3, before=14, after=6, line=LINE_TIGHT)
    h3.paragraph_format.keep_with_next = True
    _outline(h3, 2)

    # Heading 4 — Notion heading_3
    h4 = doc.styles[S.STYLE_H4]
    h4.base_style = normal
    _set_font(h4, FONT_BODY, FONT_BODY_FALLBACK, SIZE_H4, bold=True, color=INK)
    _spacing(h4, before=12, after=5, line=LINE_TIGHT)
    h4.paragraph_format.left_indent = Cm(0.25)
    _outline(h4, 3)

    # Heading 5 — Notion heading_4
    h5 = doc.styles[S.STYLE_H5]
    h5.base_style = normal
    _set_font(h5, FONT_BODY, FONT_BODY_FALLBACK, SIZE_H5, bold=True, color=INK)
    _spacing(h5, before=10, after=4, line=LINE_TIGHT)
    h5.paragraph_format.left_indent = Cm(0.5)
    _outline(h5, 4)

    # Heading 6 — rubriques fiches jurisprudence (Objet, Portée, Considérant…)
    h6 = _get_or_add(doc, S.STYLE_H6, WD_STYLE_TYPE.PARAGRAPH)
    h6.base_style = normal
    _set_font(h6, FONT_BODY, FONT_BODY_FALLBACK, SIZE_H6, bold=True, color=INK)
    _spacing(h6, before=8, after=4, line=LINE_TIGHT)
    _align(h6, WD_ALIGN_PARAGRAPH.LEFT)
    _outline(h6, None)

    # Quote — citations + considérants (.legal-prose blockquote)
    quote = _get_or_add(doc, S.STYLE_QUOTE, WD_STYLE_TYPE.PARAGRAPH)
    quote.base_style = normal
    _set_font(quote, FONT_BODY, FONT_BODY_FALLBACK, SIZE_BODY - 0.5, italic=True, color=INK)
    _spacing(quote, before=10, after=10, line=LINE_BODY)
    quote.paragraph_format.left_indent = Cm(0.5)
    quote.paragraph_format.right_indent = Cm(0.3)
    _left_border(quote, color_hex="C4A35A", sz=16, space=6)
    _shading(quote, QUOTE_BG)
    _align(quote, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _outline(quote, None)

    # Encadré — callouts Notion : corps Normal + barre bleue à gauche (marge corps)
    callout = _get_or_add(doc, S.STYLE_CALLOUT, WD_STYLE_TYPE.PARAGRAPH)
    callout.base_style = normal
    _set_font(callout, FONT_BODY, FONT_BODY_FALLBACK, SIZE_BODY, color=INK)
    _spacing(callout, before=0, after=6, line=LINE_BODY)
    _theme_left_border(callout)
    _align(callout, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _outline(callout, None)

    # List Paragraph
    try:
        list_p = doc.styles[S.STYLE_LIST_BULLET]
    except KeyError:
        list_p = _get_or_add(doc, S.STYLE_LIST_BULLET, WD_STYLE_TYPE.PARAGRAPH)
    list_p.base_style = normal
    _set_font(list_p, FONT_BODY, FONT_BODY_FALLBACK, SIZE_BODY, color=INK)
    _spacing(list_p, before=0, after=4, line=LINE_BODY)

    # List Number
    list_n = doc.styles[S.STYLE_LIST_NUMBER]
    list_n.base_style = normal
    _set_font(list_n, FONT_BODY, FONT_BODY_FALLBACK, SIZE_BODY, color=INK)
    _spacing(list_n, before=0, after=4, line=LINE_BODY)

    # Hyperlien — couleur accent (site a { color: var(--accent) })
    try:
        hyperlink = doc.styles["Hyperlink"]
        _set_font(hyperlink, FONT_BODY, FONT_BODY_FALLBACK, SIZE_BODY, color=ACCENT)
    except KeyError:
        pass

    # Table Grid — bordures simples (.legal-table)
    try:
        tbl = doc.styles["Table Grid"]
        tbl.font.name = FONT_BODY
        tbl.font.size = Pt(SIZE_BODY - 0.5)
    except KeyError:
        pass


def _set_doc_language_fr(doc: Document) -> None:
    styles_el = doc.styles.element
    for rPr in styles_el.iter(qn("w:rPr")):
        for lang in list(rPr.findall(qn("w:lang"))):
            lang.set(qn("w:val"), "fr-FR")


def _add_demo_page(doc: Document) -> None:
    """Galerie de démonstration des styles."""
    doc.add_paragraph("Titre de la page Notion", style=S.STYLE_H1)
    doc.add_paragraph(
        "Jurisprudence · Index — références relationnelles (Subtitle)",
        style=S.STYLE_SUBTITLE,
    )
    doc.add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod "
        "tempor incididunt ut labore et dolore magna aliqua.",
        style=S.STYLE_NORMAL,
    )
    doc.add_paragraph("Section Notion H1", style=S.STYLE_H2)
    doc.add_paragraph("Corps de texte justifié, police DM Sans, interligne 1,65.", style=S.STYLE_NORMAL)
    doc.add_paragraph("Section Notion H2", style=S.STYLE_H3)
    doc.add_paragraph("« Citation ou considérant de principe — barre or, fond ivoire. »", style=S.STYLE_QUOTE)
    doc.add_paragraph(
        "Encadré callout — barre bleue à gauche, texte Normal (sans fond).",
        style=S.STYLE_CALLOUT,
    )
    doc.add_paragraph("Puce un", style=S.STYLE_LIST_BULLET)
    doc.add_paragraph("Puce deux", style=S.STYLE_LIST_BULLET)
    doc.add_paragraph("Point numéroté", style=S.STYLE_LIST_NUMBER)


def _save_as_dotx(doc: Document, path: Path) -> None:
    tmp = path.with_suffix(".docx")
    doc.save(tmp)
    old = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
    new = "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
    with zipfile.ZipFile(tmp, "r") as zin, zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.decode("utf-8").replace(old, new).encode("utf-8")
            zout.writestr(item, data)
    tmp.unlink(missing_ok=True)


def build_template(*, demo: bool = True) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    _configure_styles(doc)
    _set_doc_language_fr(doc)
    if demo:
        _add_demo_page(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    _save_as_dotx(doc, OUT)
    return OUT


def main() -> None:
    path = build_template()
    print(f"Modèle : {path}")


if __name__ == "__main__":
    main()
