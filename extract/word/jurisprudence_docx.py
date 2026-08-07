"""Convertisseur fiches Jurisprudence → Word (propriétés Notion)."""

from __future__ import annotations

import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.table import _Cell

from packages.ep_core.notion import page_title, property_plain

from .converter import PageConverter
from .frames import style_fiche_decision_table
from .relations import render_relations_footer
from .styles import (
    STYLE_H6,
    STYLE_NORMAL,
    STYLE_QUOTE,
    page_title_style,
)
from .theme import ACCENT, CALLOUT_PARA_SPACE_AFTER_PT

_DATE_PREFIX = re.compile(r"^(\d{4}-\d{2}-\d{2})")

FICHE_FIELDS = (
    "Faits",
    "Enjeu juridique",
    "Solution",
    "Perspective",
)

CONSIDERANT_NAMES = ("Considérant de principe", "Considerant de principe")


def _norm(name: str) -> str:
    return (name or "").strip().casefold()


def _prop_by_name(props: dict[str, Any], *names: str) -> tuple[str, Any] | None:
    lower = {_norm(k): k for k in props}
    for name in names:
        key = lower.get(_norm(name))
        if key:
            return key, props[key]
    return None


def _prop_text(props: dict[str, Any], *names: str) -> str:
    found = _prop_by_name(props, *names)
    if not found:
        return ""
    raw = found[1]
    if isinstance(raw, str):
        return raw.strip()
    if isinstance(raw, dict) and "type" in raw:
        return property_plain(raw).strip()
    if raw is None:
        return ""
    return str(raw).strip()


def _format_date(raw: str) -> str:
    m = _DATE_PREFIX.match((raw or "").strip())
    if m:
        y, mo, d = m.group(1).split("-")
        return f"{d}/{mo}/{y}"
    return (raw or "").strip()


def _format_reference(raw: str) -> str:
    raw = (raw or "").strip()
    if not raw:
        return ""
    if raw.lower().startswith("n"):
        return raw
    if re.fullmatch(r"\d+", raw):
        return f"n° {raw}"
    return raw


def _title_from_props(props: dict[str, Any], page: dict[str, Any]) -> str:
    text = _prop_text(props, "Nom")
    if text:
        return text
    return page_title(page)


def _meta_line(props: dict[str, Any]) -> str:
    bits: list[str] = []
    for names in (
        ("Juridiction", "Jurisdiction"),
        ("Formation de jugement", "Formation"),
        ("Date",),
        ("Référence", "Reference"),
    ):
        text = _prop_text(props, *names)
        if not text:
            continue
        if _norm(names[0]) == "date":
            text = _format_date(text.split("→", 1)[0].strip())
        elif _norm(names[0]) == "référence":
            text = _format_reference(text)
        bits.append(text)
    return ", ".join(bits)


def _style_or(doc: Document, name: str, fallback: str = "Normal") -> str:
    try:
        doc.styles[name]
        return name
    except KeyError:
        return fallback


def _add_body_paragraph(doc: Document, text: str) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph(style=_style_or(doc, STYLE_NORMAL))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def _add_objet_bold(doc: Document, text: str) -> None:
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph(style=_style_or(doc, STYLE_NORMAL))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = True


def _add_heading6_label(doc: Document, label: str, *, accent: bool = False) -> None:
    p = doc.add_paragraph(label, style=_style_or(doc, STYLE_H6))
    if accent:
        for run in p.runs:
            run.font.color.rgb = ACCENT


def _host_paragraph(cell: _Cell, doc: Document, style: str | None = None) -> Any:
    if len(cell.paragraphs) == 1 and not (cell.paragraphs[0].text or "").strip():
        p = cell.paragraphs[0]
        if style:
            try:
                p.style = _style_or(doc, style)
            except KeyError:
                pass
        return p
    return cell.add_paragraph(style=_style_or(doc, style) if style else None)


def _render_fiche_box(doc: Document, props: dict[str, Any]) -> None:
    meta = _meta_line(props)
    items = [(key, _prop_text(props, key)) for key in FICHE_FIELDS]
    items = [(k, t) for k, t in items if t]
    if not items and not meta:
        return

    table = doc.add_table(rows=1, cols=1)
    style_fiche_decision_table(table)
    cell = table.rows[0].cells[0]
    cell.text = ""

    banner = _host_paragraph(cell, doc, STYLE_NORMAL)
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run("Fiche de décision")
    run.bold = True
    banner.paragraph_format.space_after = Pt(CALLOUT_PARA_SPACE_AFTER_PT)

    if meta:
        mp = _host_paragraph(cell, doc, STYLE_NORMAL)
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = mp.add_run(meta)
        meta_run.italic = True
        meta_run.font.color.rgb = ACCENT
        mp.paragraph_format.space_after = Pt(CALLOUT_PARA_SPACE_AFTER_PT)

    for label, text in items:
        p = _host_paragraph(cell, doc, STYLE_NORMAL)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        head = p.add_run(f"{label}. ")
        head.bold = True
        p.add_run(text)
        p.paragraph_format.space_after = Pt(CALLOUT_PARA_SPACE_AFTER_PT)


class JurisprudenceConverter(PageConverter):
    def render_into(
        self,
        doc: Document,
        page: dict[str, Any],
        blocks: list[dict[str, Any]],
        *,
        heading_numbering=None,
    ) -> None:
        props = page.get("properties") or {}
        doc.add_paragraph(_title_from_props(props, page), style=page_title_style())

        _add_objet_bold(doc, _prop_text(props, "Objet"))
        _add_body_paragraph(doc, _prop_text(props, "Portée"))

        has_considérant = False
        found = _prop_by_name(props, *CONSIDERANT_NAMES)
        if found:
            text = _prop_text(props, found[0])
            if text:
                has_considérant = True
                _add_heading6_label(doc, found[0], accent=True)
                q = doc.add_paragraph(text, style=_style_or(doc, STYLE_QUOTE))
                q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        has_fiche = bool(_meta_line(props)) or any(
            _prop_text(props, k) for k in FICHE_FIELDS
        )
        if has_considérant and has_fiche:
            blank = doc.add_paragraph(style=_style_or(doc, STYLE_NORMAL))
            blank.paragraph_format.space_before = Pt(0)
            blank.paragraph_format.space_after = Pt(0)

        _render_fiche_box(doc, props)

        render_relations_footer(doc, props, resolve_title=self.resolve_title)

    @classmethod
    def from_json_page(cls, page: dict[str, Any]) -> dict[str, Any]:
        """Réinjecte un enregistrement jurisprudence.json comme page API-like."""
        props_raw: dict[str, Any] = {}
        for name, value in (page.get("properties") or {}).items():
            if isinstance(value, str):
                props_raw[name] = {
                    "type": "rich_text",
                    "rich_text": [
                        {
                            "type": "text",
                            "plain_text": value,
                            "text": {"content": value},
                        }
                    ],
                }
            elif isinstance(value, dict) and "start" in value:
                props_raw[name] = {"type": "date", "date": value}
            elif isinstance(value, list):
                props_raw[name] = {
                    "type": "multi_select",
                    "multi_select": [{"name": str(x)} for x in value],
                }
            else:
                props_raw[name] = {
                    "type": "rich_text",
                    "rich_text": [
                        {
                            "type": "text",
                            "plain_text": str(value or ""),
                            "text": {"content": str(value or "")},
                        }
                    ],
                }
        title = page.get("title") or _prop_text(
            {k: v for k, v in (page.get("properties") or {}).items()}, "Nom"
        )
        if not title:
            title = "sans-titre"
        props_raw.setdefault(
            "Nom",
            {
                "type": "title",
                "title": [
                    {
                        "type": "text",
                        "plain_text": title,
                        "text": {"content": title},
                    }
                ],
            },
        )
        return {"id": page.get("id"), "properties": props_raw}
