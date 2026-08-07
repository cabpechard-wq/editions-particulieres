"""Rendu des blocs Notion → paragraphes Word."""

from __future__ import annotations

from typing import Any, Callable

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt

from .frames import finalize_callout_frames
from .heading_numbering import MAX_NUMBERED_HEADING_LEVEL, HeadingNumbering, prefix_rich_text
from .rich_text import append_rich_text, rich_plain
from .styles import (
    NOTION_BLOCK_STYLE,
    STYLE_NORMAL,
    STYLE_QUOTE,
    notion_heading_style,
)


def _style_or(doc: Document, name: str, fallback: str = "Normal") -> str:
    try:
        doc.styles[name]
        return name
    except KeyError:
        return fallback


def block_rich_text(block: dict[str, Any]) -> list[dict[str, Any]]:
    btype = block.get("type")
    data = block.get(btype) or {}
    if isinstance(data, dict):
        return data.get("rich_text") or []
    return []


def is_empty_block(block: dict[str, Any]) -> bool:
    btype = block.get("type")
    if btype == "divider":
        return False
    if btype == "table_row":
        cells = (block.get("table_row") or {}).get("cells") or []
        return not cells
    children = block.get("_children") or []
    if rich_plain(block_rich_text(block)):
        return False
    if not children:
        return True
    return all(is_empty_block(c) for c in children)


class BlockRenderer:
    def __init__(
        self,
        *,
        resolve_title: Callable[[str], str] | None = None,
        in_callout: bool = False,
        callout_paragraphs: list | None = None,
        heading_numbering: HeadingNumbering | None = None,
    ):
        self.resolve_title = resolve_title
        self.in_callout = in_callout
        self._callout_paragraphs = callout_paragraphs if callout_paragraphs is not None else []
        self.heading_numbering = heading_numbering
        self._numbered = 0

    def _child_renderer(self, **kwargs: Any) -> BlockRenderer:
        return BlockRenderer(
            resolve_title=self.resolve_title,
            in_callout=kwargs.get("in_callout", self.in_callout),
            callout_paragraphs=kwargs.get("callout_paragraphs"),
            heading_numbering=self.heading_numbering,
        )

    def render_blocks(self, doc: Document, blocks: list[dict[str, Any]]) -> None:
        for block in blocks:
            if is_empty_block(block):
                continue
            btype = block.get("type")
            if btype != "numbered_list_item":
                self._numbered = 0
            self._render_block(doc, block)

    def _stage_callout_paragraph(self, paragraph) -> None:
        if not self.in_callout:
            return
        if not (paragraph.text or "").strip():
            return
        self._callout_paragraphs.append(paragraph)

    def _finish_paragraph(self, paragraph) -> None:
        self._stage_callout_paragraph(paragraph)

    def _render_block(self, doc: Document, block: dict[str, Any]) -> None:
        btype = block.get("type")
        data = block.get(btype) or {}
        children = block.get("_children") or []

        if btype == "column_list":
            for col in children:
                if col.get("type") == "column":
                    self.render_blocks(doc, col.get("_children") or [])
            return
        if btype == "column":
            self.render_blocks(doc, children)
            return

        if btype == "callout":
            rich = data.get("rich_text") or []
            collector: list = []
            if rich_plain(rich):
                p = doc.add_paragraph(style=_style_or(doc, STYLE_NORMAL))
                append_rich_text(
                    p, rich, resolve_title=self.resolve_title
                )
                if (p.text or "").strip():
                    collector.append(p)
            nested = self._child_renderer(in_callout=True, callout_paragraphs=collector)
            nested.render_blocks(doc, children)
            finalize_callout_frames(collector)
            return

        style = NOTION_BLOCK_STYLE.get(btype, STYLE_NORMAL)

        if btype in {"heading_1", "heading_2", "heading_3", "heading_4"}:
            level = int(btype.rsplit("_", 1)[-1])
            style = notion_heading_style(level)
            rich = data.get("rich_text") or []
            if self.heading_numbering is not None and level <= MAX_NUMBERED_HEADING_LEVEL:
                rich = prefix_rich_text(rich, self.heading_numbering.next(level))
            p = doc.add_paragraph(style=_style_or(doc, style))
            append_rich_text(p, rich, resolve_title=self.resolve_title)
            self._finish_paragraph(p)
            if children:
                self.render_blocks(doc, children)
            return

        if btype == "quote":
            p = doc.add_paragraph(style=_style_or(doc, STYLE_QUOTE))
            append_rich_text(
                p, data.get("rich_text") or [], resolve_title=self.resolve_title
            )
            self._finish_paragraph(p)
            for child in children:
                if child.get("type") == "paragraph":
                    cp = doc.add_paragraph(style=_style_or(doc, STYLE_QUOTE))
                    append_rich_text(
                        cp,
                        (child.get("paragraph") or {}).get("rich_text") or [],
                        resolve_title=self.resolve_title,
                    )
                    self._finish_paragraph(cp)
                else:
                    self._render_block(doc, child)
            return

        if btype == "bulleted_list_item":
            p = doc.add_paragraph(style=_style_or(doc, style))
            append_rich_text(
                p, data.get("rich_text") or [], resolve_title=self.resolve_title
            )
            self._finish_paragraph(p)
            if children:
                self.render_blocks(doc, children)
            return

        if btype == "numbered_list_item":
            self._numbered += 1
            p = doc.add_paragraph(style=_style_or(doc, style))
            append_rich_text(
                p, data.get("rich_text") or [], resolve_title=self.resolve_title
            )
            self._finish_paragraph(p)
            if children:
                self.render_blocks(doc, children)
            return

        if btype == "to_do":
            checked = "☑" if data.get("checked") else "☐"
            p = doc.add_paragraph(style=_style_or(doc, style))
            p.add_run(f"{checked} ")
            append_rich_text(
                p, data.get("rich_text") or [], resolve_title=self.resolve_title
            )
            self._finish_paragraph(p)
            return

        if btype == "toggle":
            if rich_plain(data.get("rich_text")):
                p = doc.add_paragraph(style=_style_or(doc, STYLE_NORMAL))
                append_rich_text(
                    p, data.get("rich_text") or [], resolve_title=self.resolve_title
                )
                self._finish_paragraph(p)
            if children:
                self.render_blocks(doc, children)
            return

        if btype == "divider":
            p = doc.add_paragraph(style=_style_or(doc, STYLE_NORMAL))
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pPr.append(
                parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="C4C4C4"/>'
                    f"</w:pBdr>"
                )
            )
            return

        if btype == "table":
            self._render_table(doc, children)
            return

        if btype == "synced_block":
            self.render_blocks(doc, children)
            return

        if btype == "paragraph" or rich_plain(data.get("rich_text")):
            p = doc.add_paragraph(style=_style_or(doc, style))
            append_rich_text(
                p, data.get("rich_text") or [], resolve_title=self.resolve_title
            )
            self._finish_paragraph(p)
            if children:
                self.render_blocks(doc, children)
            return

        if children:
            self.render_blocks(doc, children)

    def _render_table(
        self, doc: Document, children: list[dict[str, Any]]
    ) -> None:
        rows_data: list[list[list[dict[str, Any]]]] = []
        for row in children:
            if row.get("type") != "table_row":
                continue
            cells = (row.get("table_row") or {}).get("cells") or []
            rows_data.append(cells)
        if not rows_data:
            return
        n_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=n_cols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        for r_idx, row in enumerate(rows_data):
            for c_idx in range(n_cols):
                cell = table.rows[r_idx].cells[c_idx]
                rich = row[c_idx] if c_idx < len(row) else []
                if rich_plain(rich):
                    p = cell.paragraphs[0]
                    append_rich_text(p, rich, resolve_title=self.resolve_title)
