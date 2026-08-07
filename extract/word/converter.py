"""Convertisseur page Notion → .docx (styles builtin V1)."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK

from packages.ep_core.notion import NotionFetcher, page_title

from .blocks import BlockRenderer
from .docx_template import open_from_template
from .heading_numbering import HeadingNumbering
from .relations import render_relations_footer
from .styles import page_title_style, template_path


class PageConverter:
    def __init__(
        self,
        template: Path | None = None,
        fetcher: NotionFetcher | None = None,
    ):
        self.template_path = Path(template or template_path())
        self.fetcher = fetcher
        self._title_cache: dict[str, str] = {}

    def resolve_title(self, page_id: str) -> str:
        if page_id in self._title_cache:
            return self._title_cache[page_id]
        title = page_id.replace("-", "")[:8]
        if self.fetcher is not None:
            try:
                page = self.fetcher.get_page(page_id)
                title = page_title(page) or title
            except Exception:
                pass
        self._title_cache[page_id] = title
        return title

    def convert(
        self,
        page: dict[str, Any],
        blocks: list[dict[str, Any]],
        out_path: Path,
    ) -> Path:
        return self.convert_many([(page, blocks)], out_path)

    def convert_many(
        self,
        pages: list[tuple[dict[str, Any], list[dict[str, Any]]]],
        out_path: Path,
        *,
        page_break: bool = True,
    ) -> Path:
        doc = open_from_template(self.template_path)
        numbering = HeadingNumbering()
        for i, (page, blocks) in enumerate(pages):
            if i > 0 and page_break:
                p = doc.add_paragraph()
                p.add_run().add_break(WD_BREAK.PAGE)
            self.render_into(doc, page, blocks, heading_numbering=numbering)
        out_path = Path(out_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return out_path

    def render_into(
        self,
        doc: Document,
        page: dict[str, Any],
        blocks: list[dict[str, Any]],
        *,
        heading_numbering: HeadingNumbering | None = None,
    ) -> None:
        props = page.get("properties") or {}
        title = page_title(page)
        doc.add_paragraph(title, style=page_title_style())
        numbering = heading_numbering if heading_numbering is not None else HeadingNumbering()
        BlockRenderer(
            resolve_title=self.resolve_title,
            heading_numbering=numbering,
        ).render_blocks(doc, blocks)
        render_relations_footer(doc, props, resolve_title=self.resolve_title)
