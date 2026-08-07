"""Vérification ponctuelle d'un modèle .dotx — usage CLI interne."""

from __future__ import annotations

import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document

REQUIRED = [
    "Title",
    "Subtitle",
    "Heading 1",
    "Heading 2",
    "Heading 3",
    "Heading 4",
    "Heading 5",
    "Normal",
    "List Paragraph",
    "List Number",
    "Quote",
    "Table Grid",
]

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def dotx_to_doc(path: Path) -> Path:
    tmp = path.with_suffix(".verify.docx")
    shutil.copy2(path, tmp)
    with zipfile.ZipFile(tmp, "r") as zin:
        ct = zin.read("[Content_Types].xml").decode("utf-8")
        ct = ct.replace("template.main+xml", "document.main+xml")
        files = {name: zin.read(name) for name in zin.namelist()}
    files["[Content_Types].xml"] = ct.encode("utf-8")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    return tmp


def outline_level(style) -> str | None:
    pPr = style.element.pPr
    if pPr is None:
        return None
    for ol in pPr.iter():
        if ol.tag == f"{{{NS_W}}}outlineLvl":
            return ol.get(f"{{{NS_W}}}val")
    return None


def verify(path: Path) -> int:
    tmp = dotx_to_doc(path)
    try:
        doc = Document(tmp)
        names = {s.name for s in doc.styles}
        missing = [n for n in REQUIRED if n not in names]

        print(f"Fichier : {path}")
        print("\n=== Styles requises ===")
        for name in REQUIRED:
            mark = "OK" if name in names else "MANQUANT"
            print(f"  [{mark}] {name}")

        print("\n=== Titres (plan + typo) ===")
        for i in range(1, 6):
            name = f"Heading {i}"
            if name not in names:
                continue
            s = doc.styles[name]
            size = s.font.size.pt if s.font.size else "?"
            print(
                f"  {name}: outline={outline_level(s)}, "
                f"font={s.font.name!r}, {size} pt, bold={s.font.bold}"
            )

        for name in ("Subtitle", "Quote", "Normal", "List Paragraph", "List Number"):
            if name not in names:
                continue
            s = doc.styles[name]
            size = s.font.size.pt if s.font.size else "?"
            pf = s.paragraph_format
            print(
                f"\n{name}: font={s.font.name!r}, {size} pt, "
                f"italic={s.font.italic}, line={pf.line_spacing}"
            )

        if "Hyperlink" in names:
            print(f"\nHyperlink: color={doc.styles['Hyperlink'].font.color.rgb}")
        else:
            styles_xml = zipfile.ZipFile(path).read("word/styles.xml").decode("utf-8")
            has_hl = "Hyperlink" in styles_xml or "Lien hypertexte" in styles_xml
            print(f"\nHyperlink (styles.xml): {'present' if has_hl else 'absent — liens utiliseront le défaut Word'}")

        sec = doc.sections[0]
        print(
            f"\nMarges (cm): L={sec.left_margin.cm:.2f} "
            f"R={sec.right_margin.cm:.2f} "
            f"T={sec.top_margin.cm:.2f} "
            f"B={sec.bottom_margin.cm:.2f}"
        )

        used = sorted({p.style.name for p in doc.paragraphs if p.style and p.style.name})
        print(f"\nCorps: {len(doc.paragraphs)} paragraphe(s)")
        if used:
            print("Styles utilisés:", ", ".join(used))

        if missing:
            print(f"\nVERDICT: ECHEC — {len(missing)} style(s) manquant(s)")
            return 1
        print("\nVERDICT: OK — compatible export Manuel V1")
        return 0
    finally:
        tmp.unlink(missing_ok=True)


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(
        "extract/templates/Editions_Particulieres-2.dotx"
    )
    raise SystemExit(verify(target))
