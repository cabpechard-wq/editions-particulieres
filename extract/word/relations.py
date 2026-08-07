"""Pied de page relations — Jurisprudence + Index (contrat V1)."""

from __future__ import annotations

from typing import Any, Callable

from docx import Document

from packages.ep_core.notion.ids import page_url_from_id

from .rich_text import add_hyperlink
from .styles import FOOTER_RELATION_KEYS, STYLE_H2, STYLE_NORMAL


def _style_or(doc: Document, name: str, fallback: str = "Normal") -> str:
    try:
        doc.styles[name]
        return name
    except KeyError:
        try:
            doc.styles[fallback]
            return fallback
        except KeyError:
            return "Normal"


def _norm(name: str) -> str:
    return (name or "").strip().casefold()


def _relation_ids(prop: dict[str, Any]) -> list[str]:
    if prop.get("type") != "relation":
        return []
    return [r["id"] for r in (prop.get("relation") or []) if r.get("id")]


def collect_relation_props(props: dict[str, Any]) -> list[tuple[str, list[str]]]:
    out: list[tuple[str, list[str]]] = []
    for name, prop in props.items():
        key = _norm(name)
        if key not in FOOTER_RELATION_KEYS:
            continue
        ids = _relation_ids(prop)
        if ids:
            out.append((name, ids))
    return out


def render_relations_footer(
    doc: Document,
    props: dict[str, Any],
    *,
    resolve_title: Callable[[str], str],
) -> None:
    """Liens en fin de document — un paragraphe par registre (Jurisprudence, Index)."""
    items = collect_relation_props(props)
    if not items:
        return

    doc.add_paragraph(
        "Ressources complémentaires",
        style=_style_or(doc, STYLE_H2),
    )

    for label, ids in items:
        p = doc.add_paragraph(style=_style_or(doc, STYLE_NORMAL))
        p.add_run(f"{label} : ")
        for j, pid in enumerate(ids):
            if j:
                p.add_run("  ·  ")
            title = resolve_title(pid)
            color = "448361" if _norm(label) == "jurisprudence" else "337EA9"
            add_hyperlink(p, title, page_url_from_id(pid), color_hex=color)
