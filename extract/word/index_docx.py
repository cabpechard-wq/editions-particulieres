"""Convertisseur entrées Glossaire (Index) → Word."""

from __future__ import annotations

from typing import Any

from docx import Document

from packages.ep_core.notion import page_title, property_plain

from .blocks import BlockRenderer
from .converter import PageConverter
from .heading_numbering import HeadingNumbering
from .relations import render_relations_footer
from .styles import STYLE_NORMAL, page_title_style


def _norm(name: str) -> str:
    return (name or "").strip().casefold()


def _prop_text(props: dict[str, Any], *names: str) -> str:
    for name, prop in props.items():
        if _norm(name) in names:
            return property_plain(prop).strip()
    return ""


class IndexConverter(PageConverter):
    def render_into(
        self,
        doc: Document,
        page: dict[str, Any],
        blocks: list[dict[str, Any]],
        *,
        heading_numbering: HeadingNumbering | None = None,
    ) -> None:
        props = page.get("properties") or {}
        doc.add_paragraph(page_title(page), style=page_title_style())

        definition = _prop_text(props, "définition", "definition")
        if definition:
            for line in definition.splitlines():
                line = line.strip()
                if line:
                    doc.add_paragraph(line, style=STYLE_NORMAL)

        source = _prop_text(props, "auteur(s) / source(s)", "auteurs / sources")
        if source:
            p = doc.add_paragraph(style=STYLE_NORMAL)
            p.add_run(source).italic = True

        if blocks:
            numbering = (
                heading_numbering if heading_numbering is not None else HeadingNumbering()
            )
            BlockRenderer(
                resolve_title=self.resolve_title,
                heading_numbering=numbering,
            ).render_blocks(doc, blocks)

        render_relations_footer(doc, props, resolve_title=self.resolve_title)
